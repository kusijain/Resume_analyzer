import streamlit as st
import docx
import pdfplumber

st.title("Resume Analyzer")
uploaded_file = st.file_uploader("Upload your resume",type = ['pdf','docx'])

def extract_text_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text+=page_text
    return text

def extract_text_from_docx(file) -> str:
    document = docx.Document(file)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)

def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError("Unsupported file type")

def clean_text(text):
    lines = text.split("\n")
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    return "\n".join(cleaned_lines)


import os
import json
from dotenv import load_dotenv
from groq import Groq

load_dotenv()
groq_key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY")
client = Groq(api_key=groq_key)

def analyze_resume(resume_text):
    prompt = f"""
You are a resume analysis assistant. Read the resume text below and return ONLY a valid JSON object with this exact structure, no extra text before or after:

{{
  "technical_skills": ["skill1", "skill2"],
  "soft_skills": ["skill1", "skill2"],
  "summary": "2-3 sentence summary of the candidate",
  "recommended_role": "one specific job role title",
  "missing_skills": ["skill1", "skill2"]
}}

Resume text:
{resume_text}
"""

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=500,
    )

    result_text = response.choices[0].message.content
    try:
        result = json.loads(result_text)
    except json.JSONDecodeError:
        result = None

    return result

if uploaded_file is not None:
    with st.spinner("Extracting text..."):
        resume_text = extract_text(uploaded_file)

    if not resume_text.strip():
        st.error("Couldn't extract any text — this may be a scanned/image-based file.")
    else:
        resume_text = clean_text(resume_text)
        with st.expander("Extracted text (clean)"):
            st.text(resume_text)

        if st.button("Analyze Resume"):
            with st.spinner("Analyzing with AI..."):
                analysis = analyze_resume(resume_text)

            if analysis is None:
                st.error("Something went wrong analyzing the resume. Please try again.")
            else:
                
                # Profile Summary
                st.subheader("📝 Profile Summary")
                st.write(analysis["summary"])

                # Recommended Role - highlighted
                st.success(f"🎯 Recommended Role: {analysis['recommended_role']}")

                # Skills side by side
                col1, col2 = st.columns(2)

                with col1:
                    st.subheader("🛠️ Technical Skills")
                    tech_tags = " ".join([f"`{skill}`" for skill in analysis["technical_skills"]])
                    st.markdown(tech_tags)

                with col2:
                    st.subheader("🤝 Soft Skills")
                    soft_tags = " ".join([f"`{skill}`" for skill in analysis["soft_skills"]])
                    st.markdown(soft_tags)

                # Missing Skills - warning box
                st.subheader("📌 Missing Skills")
                st.warning(", ".join(analysis["missing_skills"]))
